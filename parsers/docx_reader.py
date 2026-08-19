# -*- coding: utf-8 -*-
"""docx 본문을 '문단과 표가 등장한 순서 그대로' 읽는다.

기존 파이프라인은 표가 있는 docx를 표 전용 파서로만 처리해 본문 문단이 통째로
누락됐다. 여기서는 body 요소를 순회하므로 문단.표가 모두, 원래 순서대로 나온다.
"""
from docx import Document


def iter_blocks(path):
    """(kind, payload) 를 문서 순서대로 생성. kind 는 'paragraph' 또는 'table'."""
    doc = Document(path)
    body = doc.element.body
    tbl_iter = iter(doc.tables)
    par_iter = iter(doc.paragraphs)
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            p = next(par_iter, None)
            if p is None:
                continue
            txt = p.text.strip()
            if txt:
                yield 'paragraph', {'text': txt, 'style': p.style.name if p.style else None}
        elif tag == 'tbl':
            t = next(tbl_iter, None)
            if t is None:
                continue
            rows = [[c.text.strip().replace('\n', ' ') for c in r.cells] for r in t.rows]
            rows = [r for r in rows if any(r)]
            if rows:
                yield 'table', {'rows': rows}


def table_to_text(rows):
    return '\n'.join(' | '.join(r) for r in rows)


def read_docx(path):
    """문단 리스트와 표 리스트를 순서 정보(order)와 함께 돌려준다."""
    out = []
    for i, (kind, payload) in enumerate(iter_blocks(path)):
        if kind == 'paragraph':
            out.append({'order': i, 'kind': 'paragraph', 'text': payload['text'],
                        'style': payload['style']})
        else:
            out.append({'order': i, 'kind': 'table', 'text': table_to_text(payload['rows']),
                        'rows': payload['rows']})
    return out
